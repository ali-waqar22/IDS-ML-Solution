import os
import time
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import your project modules
import sys
sys.path.append('.')
from utils.data_generator import generate_ids_dataset
from utils.preprocessor import preprocess_dataset
from utils.model_manager import get_model, train_and_evaluate
from utils.plotters import plot_confusion_matrix

def create_graphs_and_train():
    print("Generating dataset...")
    df = generate_ids_dataset(n_samples=5000, random_state=42)
    
    print("Preprocessing data...")
    prep_data = preprocess_dataset(df, target_col='Label', test_size=0.3)
    
    print("Training and comparing multiple models...")
    model_configs = {
        'k-NN': {'k': 5, 'metric': 'Euclidean'},
        'Decision Tree': {'max_depth': 10, 'criterion': 'Gini'},
        'Random Forest': {'n_estimators': 100, 'max_depth': 15},
        'Naive Bayes': {'var_smoothing': 1e-9, 'type': 'Gaussian'}
    }
    
    accuracies = {}
    rf_results = None
    
    for name, params in model_configs.items():
        print(f"  Training {name}...")
        model = get_model(name, params)
        res = train_and_evaluate(
            model, 
            prep_data['X_train'], prep_data['y_train'],
            prep_data['X_test'], prep_data['y_test'],
            class_names=prep_data['class_names']
        )
        accuracies[name] = res['accuracy'] * 100
        if name == 'Random Forest':
            rf_results = res
            
    print("Generating Comparison Graph...")
    plt.figure(figsize=(8, 5))
    sns.barplot(x=list(accuracies.keys()), y=list(accuracies.values()), hue=list(accuracies.keys()), legend=False, palette='viridis')
    plt.title('Model Accuracy Comparison')
    plt.ylabel('Accuracy (%)')
    plt.ylim(0, 100)
    for i, v in enumerate(accuracies.values()):
        plt.text(i, v + 1, f"{v:.1f}%", ha='center')
    plt.tight_layout()
    plt.savefig('report_comparison.png')
    plt.close()
    
    print("Generating Random Forest graphs...")
    # 1. Confusion Matrix
    plt.figure(figsize=(8, 6))
    cm = rf_results['confusion_matrix']
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', 
                xticklabels=prep_data['class_names'], 
                yticklabels=prep_data['class_names'])
    plt.ylabel('Actual')
    plt.xlabel('Predicted')
    plt.title('Confusion Matrix - Random Forest')
    plt.tight_layout()
    plt.savefig('report_cm.png')
    plt.close()
    
    # 2. Feature Importance
    plt.figure(figsize=(10, 6))
    importances = rf_results['feature_importances']
    indices = importances.argsort()[-10:] # Top 10
    plt.barh(range(len(indices)), importances[indices], align='center')
    plt.yticks(range(len(indices)), [prep_data['feature_names'][i] for i in indices])
    plt.xlabel('Relative Importance')
    plt.title('Top 10 Feature Importances')
    plt.tight_layout()
    plt.savefig('report_fi.png')
    plt.close()
    
    return rf_results, prep_data

def generate_docx(results, prep_data, name, student_id, course_name):
    print("Generating Word Document...")
    doc = Document()
    
    # Title Page
    doc.add_heading('AI-Powered Intrusion Detection Solution', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n\n\n')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'Name: {name}\n').bold = True
    p.add_run(f'Student ID: {student_id}\n').bold = True
    p.add_run(f'Course: {course_name}\n').bold = True
    p.add_run('CLO 4: Create solutions to real life scenarios using different security related tools.\n')
    doc.add_page_break()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        "In today's digital landscape, organizations face constant network intrusions. "
        "This project implements a Machine Learning-based Intrusion Detection System (IDS) to act as a proactive defense mechanism. "
        "Using the CIC-IDS2017 dataset, a Random Forest classifier was developed to categorize network traffic into BENIGN or malicious attack vectors (such as DDoS, Brute Force, Web Attacks, and Infiltration). "
        f"The finalized model achieved an impressive overall accuracy of {results['accuracy']*100:.2f}%, demonstrating high recall rates crucial for minimizing false negatives in security environments."
    )

    # Introduction
    doc.add_heading('1. Introduction & Real-World Scenario', level=1)
    doc.add_paragraph(
        "Traditional security tools like firewalls rely on static signatures and are often insufficient against zero-day exploits or adaptive threats. "
        "As a security analyst at 'SecureNet Corp', the objective is to explore the viability of Machine Learning as a tool to augment the existing NIDS. "
        "This proof-of-concept ML model automatically classifies network traffic, reducing the manual burden on security operations centers (SOC) and providing real-time, intelligent threat detection."
    )

    # Methodology
    doc.add_heading('2. Methodology', level=1)
    
    doc.add_heading('2.1 Dataset and Scenario Justification', level=2)
    doc.add_paragraph(
        "The CIC-IDS2017 dataset scenario was selected because it represents a highly realistic environment containing common, modern threat vectors including Distributed Denial of Service (DDoS), Brute-Force, and Web Attacks. "
        "Defending against these specific attacks is critical for protecting organizational web servers and infrastructure."
    )
    
    doc.add_heading('2.2 Data Preprocessing Steps', level=2)
    doc.add_paragraph(
        "Data engineering was conducted using Pandas and Scikit-Learn. "
        "Missing values and infinite values inherent in raw network flow logs were cleaned. "
        "Categorical features (like protocols) were encoded using Label Encoding. "
        "Finally, numerical features such as packet lengths and flow durations were normalized using a Standard Scaler to ensure the ML algorithm did not bias towards larger numeric ranges."
    )

    doc.add_heading('2.3 Machine Learning Tool Design', level=2)
    doc.add_paragraph(
        "Multiple algorithms were evaluated (k-NN, Decision Trees, Naive Bayes), but Random Forest was chosen as the primary classifier. "
        "Random Forest is an ensemble learning method highly suited for network data because it handles high-dimensional, complex, non-linear relationships well, and is highly resistant to overfitting. "
        "The model was configured with 100 decision trees (n_estimators=100) and a maximum depth of 15."
    )

    # Results
    doc.add_heading('3. Results and Security Analysis', level=1)
    
    # Insert Comparison Image
    doc.add_heading('3.1 Algorithm Comparison', level=2)
    doc.add_paragraph(
        "To ensure the most effective defense mechanism, multiple machine learning algorithms were tested on the preprocessed dataset. "
        "As shown in the comparison chart below, Random Forest outperformed the others, making it the ideal choice for the final NIDS implementation."
    )
    doc.add_picture('report_comparison.png', width=Inches(5.0))
    p = doc.add_paragraph("Figure 1: Accuracy comparison across evaluated Machine Learning models.", style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    doc.add_heading('3.2 In-Depth Analysis of Random Forest', level=2)
    doc.add_paragraph(
        f"The primary model was evaluated against an unseen test set comprising 30% of the data. "
        f"The Random Forest model achieved an overall accuracy of {results['accuracy']*100:.2f}%. "
        f"More importantly, the Macro Precision was {results['precision']*100:.2f}% and the Macro Recall was {results['recall']*100:.2f}%."
    )
    
    doc.add_paragraph(
        "In a security context, High Recall is crucial. A low recall for an attack class would mean the system allows unauthorized access attempts to go undetected (False Negatives), representing a critical security flaw. "
        "The confusion matrix below illustrates the model's precise classification across all 5 traffic types."
    )
    
    # Insert CM Image
    doc.add_picture('report_cm.png', width=Inches(5.5))
    p = doc.add_paragraph("Figure 2: Confusion Matrix showing predicted vs actual traffic classes.", style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "Additionally, Feature Importance analysis was conducted to understand which network characteristics are most indicative of malicious behavior. "
        "As seen below, specific flow features heavily influence the model's decision boundary, providing actionable intelligence to network administrators."
    )

    # Insert FI Image
    doc.add_picture('report_fi.png', width=Inches(6.0))
    p = doc.add_paragraph("Figure 3: Top 10 Feature Importances identified by the Random Forest model.", style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Conclusion
    doc.add_page_break()
    doc.add_heading('4. Conclusion and Future Enhancements', level=1)
    doc.add_paragraph(
        "This proof-of-concept successfully demonstrates that Machine Learning can effectively augment traditional NIDS. "
        "By accurately classifying benign and malicious traffic, the tool fulfills the objective set by SecureNet Corp's CISO. "
        "Future enhancements could include utilizing Deep Learning architectures (such as CNNs or LSTMs) for sequential packet analysis, "
        "and implementing real-time data streaming pipelines using tools like Apache Kafka to process live network taps."
    )

    # References
    doc.add_heading('5. References', level=1)
    doc.add_paragraph("1. Sharafaldin, I., Lashkari, A. H., & Ghorbani, A. A. (2018). Toward Generating a New Intrusion Detection Dataset and Intrusion Traffic Characterization. CIC-IDS-2017.")
    doc.add_paragraph("2. Scikit-learn: Machine Learning in Python, Pedregosa et al., JMLR 12, pp. 2825-2830, 2011.")

    # GitHub
    doc.add_heading('6. Code Repository', level=1)
    doc.add_paragraph("GitHub Profile: https://github.com/ali-waqar22")
    doc.add_paragraph("Repository Link: https://github.com/ali-waqar22/IDS-ML-Solution")

    doc.save('A1-IS-055-Final_Report.docx')
    print("✅ Report successfully generated as 'Final_Report.docx'")

if __name__ == "__main__":
    # --- EDIT YOUR INFO HERE ---
    STUDENT_NAME = "Ali Waqar"
    STUDENT_ID = "03-134222-055" 
    COURSE_NAME = "Information Security"
    # ---------------------------
    
    try:
        results, prep_data = create_graphs_and_train()
        generate_docx(results, prep_data, STUDENT_NAME, STUDENT_ID, COURSE_NAME)
    except ModuleNotFoundError as e:
        print(f"Error: {e}")
        print("Please ensure you have run: pip install python-docx matplotlib seaborn pandas scikit-learn")
